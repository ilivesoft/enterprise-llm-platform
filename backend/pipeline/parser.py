from __future__ import annotations

import re
from pathlib import Path
from dataclasses import dataclass, field

from pipeline.exceptions import UnsupportedFileTypeError, EmptyFileError

SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".txt",
    ".md",
    ".html",
    ".png",
    ".jpg",
    ".jpeg",
    ".bmp",
    ".tiff",
}

_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".bmp", ".tiff"}


@dataclass
class ParsedDocument:
    text: str
    filename: str
    file_type: str
    metadata: dict = field(default_factory=dict)


# @MX:ANCHOR: [AUTO] 문서 파싱 진입점 — rag_service, ingest 엔드포인트에서 호출
# @MX:REASON: fan_in >= 3; 반환 구조(ParsedDocument) 변경 시 파이프라인 전체 영향
class DocumentParser:
    """Parses documents into plain text.

    Supported formats: PDF, DOCX, TXT, MD, HTML, PNG, JPG, JPEG, BMP, TIFF.
    OCR is applied to image files and scanned PDF pages (lazy-loaded).
    Falls back to UnstructuredFileLoader for complex PDF/DOCX/HTML documents.
    """

    def __init__(self) -> None:
        self._ocr_reader = None

    # @MX:WARN: [AUTO] EasyOCR 초기화 — 수백 MB 모델 다운로드 발생 가능
    # @MX:REASON: 첫 호출 시 모델 다운로드; 네트워크 제한 환경에서 타임아웃 위험
    def _get_ocr_reader(self):
        """Lazily initialises the EasyOCR reader (downloads model on first call)."""
        if self._ocr_reader is None:
            import easyocr
            from config.settings import settings

            self._ocr_reader = easyocr.Reader(
                settings.ocr_languages,
                gpu=settings.ocr_gpu,
            )
        return self._ocr_reader

    # @MX:ANCHOR: [AUTO] 파서 공개 API — rag_service, ingest, 테스트에서 호출
    # @MX:REASON: fan_in >= 3; 반환 타입(ParsedDocument) 변경 시 rag_service 호환성 깨짐
    def parse(self, file_path: Path) -> ParsedDocument:
        """Auto-detects file format and extracts text.

        Raises:
            FileNotFoundError: File does not exist.
            UnsupportedFileTypeError: Extension not in SUPPORTED_EXTENSIONS.
            EmptyFileError: No text could be extracted.
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        if not self.is_supported(file_path):
            raise UnsupportedFileTypeError(
                f"Unsupported file type: {file_path.suffix}"
            )

        ext = file_path.suffix.lower()
        metadata: dict = {}

        if ext == ".pdf":
            text, metadata = self._parse_pdf(file_path)
        elif ext == ".docx":
            text = self._parse_docx(file_path)
        elif ext == ".html":
            text = self._parse_html(file_path)
        elif ext in _IMAGE_EXTENSIONS:
            text = self._parse_image(file_path)
        else:
            # .txt, .md
            text = self._parse_text(file_path)

        # Fallback to UnstructuredFileLoader for complex formats
        _unstructured_exts = {".pdf", ".docx", ".html"}
        if (not text or not text.strip()) and ext in _unstructured_exts:
            try:
                text = self._parse_with_unstructured(file_path)
            except Exception:
                pass

        if not text or not text.strip():
            raise EmptyFileError(f"No text could be extracted: {file_path}")

        return ParsedDocument(
            text=text,
            filename=file_path.name,
            file_type=ext.lstrip("."),
            metadata=metadata,
        )

    def _parse_pdf(self, file_path: Path) -> tuple[str, dict]:
        """Extracts text with PyMuPDF; falls back to OCR for scanned pages."""
        import fitz  # type: ignore[import-untyped]

        text_parts: list[str] = []
        with fitz.open(str(file_path)) as doc:
            pages = len(doc)
            for page in doc:
                page_text = page.get_text()
                if page_text and page_text.strip():
                    text_parts.append(page_text)
                else:
                    # Scanned page — render to PNG bytes then OCR
                    pix = page.get_pixmap()
                    img_bytes = pix.tobytes("png")
                    ocr_text = self._ocr_from_bytes(img_bytes)
                    if ocr_text:
                        text_parts.append(ocr_text)

        return "\n".join(text_parts), {"pages": pages}

    def _parse_docx(self, file_path: Path) -> str:
        """Extracts text from DOCX paragraphs and tables (deduplicates merged cells)."""
        from docx import Document  # type: ignore[import-untyped]

        doc = Document(str(file_path))
        text_parts: list[str] = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                # Remove consecutive duplicates caused by merged cells
                deduped: list[str] = []
                for t in row_texts:
                    if not deduped or t != deduped[-1]:
                        deduped.append(t)
                if deduped:
                    text_parts.append(" | ".join(deduped))

        return "\n".join(text_parts)

    def _parse_html(self, file_path: Path) -> str:
        """Strips HTML tags and returns plain text."""
        html = file_path.read_text(encoding="utf-8", errors="replace")
        text = re.sub(r"<[^>]+>", " ", html)
        text = re.sub(r"\s+", " ", text).strip()
        return text

    def _parse_image(self, file_path: Path) -> str:
        """Applies EasyOCR to an image file.

        Reads the file as bytes to avoid cv2.imread Unicode path issues on Windows.
        """
        return self._ocr_from_bytes(file_path.read_bytes())

    def _ocr_from_bytes(self, img_bytes: bytes) -> str:
        """Runs EasyOCR on raw image bytes (PNG / JPEG / etc.)."""
        import numpy as np  # type: ignore[import-untyped]
        import cv2  # type: ignore[import-untyped]

        reader = self._get_ocr_reader()
        nparr = np.frombuffer(img_bytes, dtype=np.uint8)
        img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
        results = reader.readtext(img, detail=0)
        return "\n".join(results)

    def _parse_text(self, file_path: Path) -> str:
        """Reads plain-text or Markdown files as UTF-8."""
        return file_path.read_text(encoding="utf-8", errors="replace")

    def _parse_with_unstructured(self, file_path: Path) -> str:
        """Fallback parser using langchain-community UnstructuredFileLoader."""
        from langchain_community.document_loaders import UnstructuredFileLoader  # type: ignore[import-untyped]

        loader = UnstructuredFileLoader(str(file_path), mode="single")
        docs = loader.load()
        return "\n\n".join(
            doc.page_content for doc in docs if doc.page_content.strip()
        )

    def is_supported(self, file_path: Path) -> bool:
        """Returns True if the file extension is in SUPPORTED_EXTENSIONS."""
        return Path(file_path).suffix.lower() in SUPPORTED_EXTENSIONS
